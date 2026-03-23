import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches

def run():
    st.subheader("Liquid Limit by Cone Penetration Method (IS 2720: Part 5: 1985)")

    # --- Procedure & Formulas ---
    with st.expander("📝 Procedure & Formulas", expanded=True):
        st.markdown("""
### Test Procedure
1. Take air-dried soil and pass it through a 425 µm sieve.
2. Mix soil with water to form a uniform paste.
3. Fill the cone penetration device with soil paste.
4. Measure penetration (in mm) for several trials.
5. Record weights:
    - W1: Empty container
    - W2: Container + wet soil
    - W3: Container + dry soil
6. Calculate water content and plot flow curve (water content vs penetration).
7. Determine **Liquid Limit (LL)** at 20 mm penetration.

### Formulas Used
**Water Content (%)**  
w (%) = ((W2 - W3) / (W3 - W1)) * 100

**Linear Regression for Flow Curve**  
Water Content = a * Penetration + b

**Liquid Limit (LL)**  
LL = Water Content at Penetration = 20 mm
""")

    # --- Helper function to calculate Moisture Content ---
    def calculate_moisture_content(w1, w2, w3):
        if w1==0.0 and w2==0.0 and w3==0.0: return 0.0
        if not (w2 > w3 > w1): return np.nan
        weight_of_water = w2 - w3
        weight_of_dry_soil = w3 - w1
        if weight_of_dry_soil <= 0: return np.nan
        return (weight_of_water / weight_of_dry_soil) * 100

    # --- Number of Trials ---
    if "cone_num_trials" not in st.session_state:
        st.session_state.cone_num_trials = 4
    num_trials = st.number_input(
        "Number of Trials", min_value=3, max_value=10,
        value=st.session_state.cone_num_trials
    )

    # --- Initialize trial inputs ---
    if "cone_trial_inputs" not in st.session_state:
        st.session_state.cone_trial_inputs = [
            {"penetration":0.0,"w1":0.0,"w2":0.0,"w3":0.0,"water_content":0.0} 
            for _ in range(num_trials)
        ]
    if len(st.session_state.cone_trial_inputs) != num_trials:
        new_inputs=[]
        for i in range(num_trials):
            if i < len(st.session_state.cone_trial_inputs):
                new_inputs.append(st.session_state.cone_trial_inputs[i])
            else:
                new_inputs.append({"penetration":0.0,"w1":0.0,"w2":0.0,"w3":0.0,"water_content":0.0})
        st.session_state.cone_trial_inputs = new_inputs
        st.session_state.cone_num_trials = num_trials

    # --- Input Fields ---
    for i in range(num_trials):
        st.markdown(f"### Trial {i+1}")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.session_state.cone_trial_inputs[i]["penetration"] = st.number_input(
                f"Penetration (mm) [{i+1}]", min_value=0.0,
                value=st.session_state.cone_trial_inputs[i]["penetration"], format="%.2f"
            )
        with col2:
            st.session_state.cone_trial_inputs[i]["w1"] = st.number_input(
                f"W1 (Empty) [{i+1}]", min_value=0.0,
                value=st.session_state.cone_trial_inputs[i]["w1"], format="%.2f"
            )
        with col3:
            st.session_state.cone_trial_inputs[i]["w2"] = st.number_input(
                f"W2 (Wet + Container) [{i+1}]", min_value=0.0,
                value=st.session_state.cone_trial_inputs[i]["w2"], format="%.2f"
            )
        with col4:
            st.session_state.cone_trial_inputs[i]["w3"] = st.number_input(
                f"W3 (Dry + Container) [{i+1}]", min_value=0.0,
                value=st.session_state.cone_trial_inputs[i]["w3"], format="%.2f"
            )

        # Calculate water content
        w1_val = st.session_state.cone_trial_inputs[i]["w1"]
        w2_val = st.session_state.cone_trial_inputs[i]["w2"]
        w3_val = st.session_state.cone_trial_inputs[i]["w3"]
        calculated_wc = calculate_moisture_content(w1_val, w2_val, w3_val)
        st.session_state.cone_trial_inputs[i]["water_content"] = calculated_wc
        if np.isnan(calculated_wc):
            st.error(f"Trial {i+1} Water Content: Invalid input.")
        else:
            st.info(f"Trial {i+1} Water Content: **{calculated_wc:.2f}%**")

    # --- Save Inputs CSV ---
    if st.button("💾 Save Inputs"):
        df_save = pd.DataFrame(st.session_state.cone_trial_inputs)
        df_save.insert(0, "Trial", range(1,len(df_save)+1))
        buffer = StringIO()
        df_save.to_csv(buffer,index=False)
        buffer.seek(0)
        st.download_button(
            label="📥 Download Input Data as CSV",
            data=buffer.getvalue(),
            file_name="cone_penetration_inputs.csv",
            mime="text/csv"
        )

    # --- Calculate Liquid Limit ---
    if st.button("Calculate Liquid Limit"):
        df_all = pd.DataFrame(st.session_state.cone_trial_inputs)
        df_all.insert(0,"Trial", range(1,len(df_all)+1))
        df_valid = df_all[(df_all["penetration"]>0) & (~df_all["water_content"].isna()) & (df_all["water_content"]>0)].copy()
        if df_valid.empty or len(df_valid)<2:
            st.error("Enter at least two valid data points with non-zero penetration and calculable water content.")
            return None
        try:
            x_data = df_valid["penetration"].astype(float)
            y_data = df_valid["water_content"].astype(float)
            if len(np.unique(x_data))<2:
                st.error("At least two distinct penetration readings required.")
                return None
            coeffs = np.polyfit(x_data,y_data,1)
            poly = np.poly1d(coeffs)
            liquid_limit = poly(20) # LL at 20mm

            # Display trial data
            st.markdown("### Trial Data & Water Content")
            st.dataframe(df_all.round(2), use_container_width=True)

            # Plot flow curve
            fig,ax=plt.subplots(figsize=(8,5))
            ax.scatter(x_data,y_data,color='blue',label='Observed Data')
            x_plot = np.linspace(min(x_data)-5,max(x_data)+5,100)
            y_plot = poly(x_plot)
            ax.plot(x_plot,y_plot,color='green',linestyle='--',label='Fitted Curve')
            ax.axvline(20,color='red',linestyle=':',label='20 mm Penetration')
            ax.axhline(liquid_limit,color='orange',linestyle=':',label=f'LL = {liquid_limit:.2f}%')
            ax.plot(20,liquid_limit,'ro',markersize=8,label='Liquid Limit Point')
            ax.set_xlabel("Penetration (mm)")
            ax.set_ylabel("Water Content (%)")
            ax.set_title("Cone Penetration Method: Flow Curve")
            ax.legend(); ax.grid(True)
            st.pyplot(fig)

            img_buf = BytesIO()
            fig.savefig(img_buf,format='png',bbox_inches='tight'); img_buf.seek(0)
            plt.close(fig)

            st.markdown("### Result")
            st.success(f"Liquid Limit (Cone Penetration Method) = {liquid_limit:.2f}%")

            # Soil Classification
            st.markdown("### Soil Classification Based on Liquid Limit")
            if liquid_limit<35: soil_class="Low Plasticity Soil"; st.info(soil_class)
            elif 35<=liquid_limit<=50: soil_class="Intermediate Plasticity Soil"; st.warning(soil_class)
            else: soil_class="High Plasticity Soil"; st.success(soil_class)

            # --- Generate Word Report ---
            if st.button("📄 Generate Word Report"):
                doc = Document()
                doc.add_heading("Liquid Limit Test Report (Cone Penetration Method)",0)
                doc.add_paragraph(f"Estimated Liquid Limit (LL) = {liquid_limit:.2f}%")
                doc.add_paragraph("Remarks: Determined at 20 mm penetration from flow curve.")
                doc.add_heading("Trial Data",level=1)
                table=doc.add_table(rows=1,cols=len(df_all.columns))
                hdr_cells=table.rows[0].cells
                for idx,col in enumerate(df_all.columns): hdr_cells[idx].text=col
                for i in range(len(df_all)):
                    row_cells=table.add_row().cells
                    for j,col in enumerate(df_all.columns): row_cells[j].text=str(df_all[col].iloc[i])
                doc.add_heading("Flow Curve",level=1)
                doc.add_picture(img_buf,width=Inches(6))
                buffer_word=BytesIO(); doc.save(buffer_word); buffer_word.seek(0)
                st.download_button("📥 Download Word Report",data=buffer_word.getvalue(),
                                   file_name="Cone_Penetration_Liquid_Limit_Report.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            return {
                "Input Data": df_all,
                "Valid Data Used for Fit": df_valid,
                "Flow Curve Graph": img_buf,
                "Liquid Limit (LL)": f"{liquid_limit:.2f}%",
                "Soil Classification": soil_class,
                "Remarks": "Liquid Limit determined at 20mm penetration."
            }

        except np.linalg.LinAlgError:
            st.error("Cannot fit curve: check that penetration values are distinct.")
            return None
        except Exception as e:
            st.error(f"Unexpected error: {e}")
            return None

    return None