import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime

def run():

    st.header("🧪 Constant Head Permeability Test")
    st.caption("IS 2720 (Part 36): 1987")

    # ---------------- SESSION STATE ----------------
    if "ch_results" not in st.session_state:
        st.session_state.ch_results = None

    if "ch_avg_k" not in st.session_state:
        st.session_state.ch_avg_k = None

    if "ch_soil" not in st.session_state:
        st.session_state.ch_soil = None

    if "ch_graph" not in st.session_state:
        st.session_state.ch_graph = None

    # ---------------- PROCEDURE ----------------
    with st.expander("📘 View Detailed Procedure"):
        st.markdown("""
### Objective
To determine the coefficient of permeability (k) of coarse-grained soil.

### Formula
k = (Q × L) / (A × h × t)

### Steps
1. Prepare soil specimen.
2. Apply constant head.
3. Collect discharge.
4. Calculate permeability.
""")

    # ---------------- INPUT ----------------
    trials = st.number_input("Number of Trials", 1, 10, 3)

    data = []

    for i in range(trials):
        st.subheader(f"Trial {i+1}")
        col1, col2, col3 = st.columns(3)

        with col1:
            L = st.number_input(
                f"L (cm) - T{i+1}",
                value=0.0,
                step=0.01,
                format="%.3f",
                key=f"L{i}"
            )

            A = st.number_input(
                f"A (cm²) - T{i+1}",
                value=0.0,
                step=0.01,
                format="%.3f",
                key=f"A{i}"
            )

        with col2:
            h = st.number_input(
                f"h (cm) - T{i+1}",
                value=0.0,
                step=0.01,
                format="%.3f",
                key=f"h{i}"
            )

            Q = st.number_input(
                f"Q (cm³) - T{i+1}",
                value=0.0,
                step=0.01,
                format="%.3f",
                key=f"Q{i}"
            )

        with col3:
            t = st.number_input(
                f"t (s) - T{i+1}",
                value=0.0,
                step=0.01,
                format="%.3f",
                key=f"t{i}"
            )

        if L > 0 and A > 0 and h > 0 and Q > 0 and t > 0:
            k = (Q * L) / (A * h * t)
        else:
            k = 0

        data.append([L, A, h, Q, t, k])

    # ---------------- CALCULATE ----------------
    if st.button("📊 Calculate"):

        df = pd.DataFrame(data, columns=[
            "Length (cm)", "Area (cm²)", "Head (cm)",
            "Volume (cm³)", "Time (s)", "k (cm/s)"
        ])

        df["k (m/s)"] = df["k (cm/s)"] / 100

        avg_k = df["k (cm/s)"].mean()

        # Soil classification
        if avg_k > 1e-1:
            soil = "Gravel"
        elif avg_k > 1e-2:
            soil = "Coarse Sand"
        elif avg_k > 1e-3:
            soil = "Medium Sand"
        elif avg_k > 1e-4:
            soil = "Fine Sand"
        elif avg_k > 1e-6:
            soil = "Silt"
        else:
            soil = "Clay"

        # Graph
        fig, ax = plt.subplots()
        ax.plot(range(1, trials+1), df["k (cm/s)"], marker='o')
        ax.set_xlabel("Trial Number")
        ax.set_ylabel("k (cm/s)")
        ax.set_title("Permeability vs Trial")
        ax.grid(True)

        img_stream = BytesIO()
        fig.savefig(img_stream, format="png")
        img_stream.seek(0)

        # Save state
        st.session_state.ch_results = df
        st.session_state.ch_avg_k = avg_k
        st.session_state.ch_soil = soil
        st.session_state.ch_graph = img_stream

    # ---------------- DISPLAY RESULTS ----------------
    if st.session_state.ch_results is not None:

        st.subheader("📋 Results")
        st.dataframe(st.session_state.ch_results.style.format(precision=5))

        st.success(f"Average k = {st.session_state.ch_avg_k:.5f} cm/s")
        st.info(f"Soil Type: {st.session_state.ch_soil}")

        st.image(st.session_state.ch_graph)

        # ---------------- WORD REPORT ----------------
        if st.button("📄 Generate Word Report"):

            df = st.session_state.ch_results
            avg_k = st.session_state.ch_avg_k
            soil = st.session_state.ch_soil
            graph = st.session_state.ch_graph

            doc = Document()
            doc.add_heading("Constant Head Permeability Test Report", 0)
            doc.add_paragraph("IS 2720 (Part 36): 1987")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

            doc.add_heading("Results", 1)

            table = doc.add_table(rows=1, cols=len(df.columns))
            for i, col in enumerate(df.columns):
                table.rows[0].cells[i].text = col

            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = f"{val:.5f}"

            doc.add_paragraph(f"\nAverage k = {avg_k:.5f} cm/s")
            doc.add_paragraph(f"Soil Type: {soil}")

            doc.add_heading("Conclusion", 1)
            doc.add_paragraph(
                f"The coefficient of permeability is {avg_k:.5f} cm/s, indicating {soil} soil."
            )

            graph.seek(0)
            doc.add_picture(graph, width=Inches(5))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download Report",
                data=buffer,
                file_name="Constant_Head_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    return None