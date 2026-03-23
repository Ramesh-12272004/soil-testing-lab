import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

def run():

    st.subheader("Sieve Analysis (IS 2720 Part 4)")

    # -----------------------------
    # PROCEDURE & FORMULAS
    # -----------------------------
    procedure_text = """
🎯 Objective:
To determine the particle size distribution of coarse-grained soil using a standard set of sieves.

🧪 Apparatus Required:
- Standard IS sieves from 4.75 mm to 75 micron
- Sieve shaker
- Weighing balance
- Oven
- Brush & tray

📝 Procedure:
1. Take a representative oven-dried soil sample.
2. Record its total dry weight.
3. Arrange sieves in descending order.
4. Place soil in top sieve and shake for 10–15 minutes.
5. Weigh and record the soil retained on each sieve.
6. Compute % retained and % passing.
7. Plot grain size distribution on a semi-log graph.
"""

    formulas_text = """
📐 Formulas Used:

% Retained = (Weight Retained / Total Weight) × 100

Cumulative % Retained = Sum of % Retained

% Passing = 100 - Cumulative % Retained

Cu = D60 / D10

Cc = (D30)^2 / (D60 × D10)

Where:
- D10 = Particle size at 10% passing
- D30 = Particle size at 30% passing
- D60 = Particle size at 60% passing
"""

    st.markdown("## 📘 Test Procedure")
    st.markdown(procedure_text)
    st.markdown("## 📐 Formulas")
    st.markdown(formulas_text)

    # -----------------------------
    # INPUT WEIGHTS
    # -----------------------------
    sieve_sizes = [4.75, 2.36, 1.18, 0.6, 0.425, 0.3, 0.15, 0.075, 0.0]
    sieve_labels = [str(s) if s != 0.0 else "Pan" for s in sieve_sizes]

    if "sieve_weights" not in st.session_state:
        st.session_state.sieve_weights = [0.0] * len(sieve_sizes)

    st.markdown("## 📥 Enter Weight Retained (grams)")
    cols = st.columns(2)
    weights_input = []
    for i, label in enumerate(sieve_labels):
        with (cols[0] if i < len(sieve_labels)/2 else cols[1]):
            weight = st.number_input(
                f"Weight Retained on {label} mm",
                min_value=0.0,
                step=0.1,
                value=st.session_state.sieve_weights[i],
                key=f"sieve_input_{i}"
            )
            weights_input.append(weight)

    st.session_state.sieve_weights = weights_input

    # -----------------------------
    # CALCULATION
    # -----------------------------
    if st.button("🔍 Calculate Sieve Analysis"):
        total_weight = sum(weights_input)
        if total_weight == 0:
            st.error("Please enter valid weights.")
            return None

        percent_retained = [(w / total_weight) * 100 for w in weights_input]
        cumulative_retained = pd.Series(percent_retained).cumsum()
        percent_passing = 100 - cumulative_retained
        percent_passing.iloc[-1] = 0

        results_df = pd.DataFrame({
            "Sieve Size (mm)": sieve_labels,
            "Weight Retained (g)": weights_input,
            "% Retained": percent_retained,
            "Cumulative % Retained": cumulative_retained,
            "% Passing": percent_passing
        })

        # Rounded table for display
        results_display = results_df.copy()
        results_display[["Weight Retained (g)", "% Retained", "Cumulative % Retained", "% Passing"]] = \
            results_display[["Weight Retained (g)", "% Retained", "Cumulative % Retained", "% Passing"]].round(2)

        st.success("Calculation Completed")
        st.dataframe(results_display, use_container_width=True)

        # -------- Grain Size Curve --------
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.semilogx(sieve_sizes[:-1], percent_passing[:-1], marker="o")
        ax.set_xlabel("Sieve Size (mm)")
        ax.set_ylabel("% Finer")
        ax.set_title("Grain Size Distribution Curve")
        ax.grid(True, which="both")
        ax.invert_xaxis()
        ax.set_ylim(0, 100)
        st.pyplot(fig)

        graph_buffer = BytesIO()
        fig.savefig(graph_buffer, format="PNG")
        graph_buffer.seek(0)
        plt.close(fig)

        # -------- D10, D30, D60, Cu & Cc --------
        def interpolate(p, s, target):
            for j in range(1, len(p)):
                if p[j-1] >= target >= p[j]:
                    x1, y1 = s[j-1], p[j-1]
                    x2, y2 = s[j], p[j]
                    return x1 + (target - y1)*(x2 - x1)/(y2 - y1)
            return None

        D10 = interpolate(list(percent_passing), sieve_sizes, 10)
        D30 = interpolate(list(percent_passing), sieve_sizes, 30)
        D60 = interpolate(list(percent_passing), sieve_sizes, 60)
        Cu = D60 / D10 if D10 else None
        Cc = (D30**2) / (D60 * D10) if D10 and D60 else None

        st.markdown("### 📊 Soil Properties")
        st.write(f"D10 = {D10:.3f} mm")
        st.write(f"D30 = {D30:.3f} mm")
        st.write(f"D60 = {D60:.3f} mm")
        st.write(f"Coefficient of Uniformity (Cu) = {Cu:.2f}")
        st.write(f"Coefficient of Curvature (Cc) = {Cc:.2f}")

        conclusion = "Soil is well graded." if Cu >= 5 else "Soil is poorly graded."
        st.markdown("### ✅ Conclusion")
        st.write(conclusion)

        # -------- Suggestion based on IS 2720 --------
        st.markdown("### 💡 Suggested Use")
        if Cu >= 5:
            st.write("According to IS 2720, this well-graded soil is suitable for road base, sub-base, and general foundation works.")
        else:
            st.write("According to IS 2720, this poorly graded soil is more suitable for embankments or fill material but may require stabilization for structural foundations.")

        # -------- WORD REPORT --------
        doc = Document()
        doc.add_heading("Soil Sieve Analysis Report", 0)
        # Procedure
        doc.add_heading("Test Procedure", level=1)
        for ln in procedure_text.strip().split("\n"):
            doc.add_paragraph(ln.strip())
        # Formulas
        doc.add_heading("Formulas", level=1)
        for ln in formulas_text.strip().split("\n"):
            doc.add_paragraph(ln.strip())
        # Results Table
        doc.add_page_break()
        doc.add_heading("Results", level=1)
        table = doc.add_table(rows=1, cols=len(results_display.columns))
        for i, col in enumerate(results_display.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in results_display.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        # Graph
        doc.add_paragraph("")
        doc.add_heading("Grain Size Distribution Curve", level=1)
        graph_buffer.seek(0)
        doc.add_picture(graph_buffer, width=Inches(5))
        # Soil Properties
        doc.add_heading("Soil Properties", level=1)
        doc.add_paragraph(f"D10 = {D10:.3f} mm")
        doc.add_paragraph(f"D30 = {D30:.3f} mm")
        doc.add_paragraph(f"D60 = {D60:.3f} mm")
        doc.add_paragraph(f"Cu = {Cu:.2f}")
        doc.add_paragraph(f"Cc = {Cc:.2f}")
        # Conclusion & Suggested Use
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(conclusion)
        doc.add_heading("Suggested Use", level=1)
        if Cu >= 5:
            doc.add_paragraph("According to IS 2720, this well-graded soil is suitable for road base, sub-base, and general foundation works.")
        else:
            doc.add_paragraph("According to IS 2720, this poorly graded soil is more suitable for embankments or fill material but may require stabilization for structural foundations.")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="⬇️ Download Word Report",
            data=buffer,
            file_name="Sieve_Analysis_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        return {"data": results_df, "graph": graph_buffer}

    return None