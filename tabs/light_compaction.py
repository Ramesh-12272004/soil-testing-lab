import streamlit as st
import matplotlib.pyplot as plt
import pandas as pd
import math
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches

def run():
    st.subheader("🧪 Light Compaction Test (IS 2720 Part 7:1980)")

    # ✅ INPUT DEFINITIONS + FORMULAS
    st.markdown("""
### 📌 Meaning of Inputs

**Required Inputs Per Trial:**

- **W1** = Weight of empty moisture content cup (g)  
- **W2** = Weight of cup + wet soil (g)  
- **W3** = Weight of cup + dry soil (g)  
- **W4** = Weight of mould + base plate (g)  
- **W5** = Weight of mould + compacted soil + base plate (g)  

---

### 📘 Key Formulas

1. **Water Content (%)**  
   w = ((W2 − W1) − (W3 − W1)) / (W3 − W1) × 100  

2. **Wet Density (g/cc)**  
   ρ₍wet₎ = (W5 − W4) / Volume  

3. **Dry Density (g/cc)**  
   ρ₍dry₎ = ρ₍wet₎ / (1 + w/100)
""")

    # --- Session State Initialization ---
    if "lc_mould_dia" not in st.session_state: st.session_state.lc_mould_dia = 10.0
    if "lc_mould_height" not in st.session_state: st.session_state.lc_mould_height = 12.7
    if "lc_num_points" not in st.session_state: st.session_state.lc_num_points = 5
    if "lc_trial_inputs" not in st.session_state:
        st.session_state.lc_trial_inputs = [{"W1":0.0,"W2":0.0,"W3":0.0,"W4":0.0,"W5":0.0} for _ in range(st.session_state.lc_num_points)]

    # Number of trials
    num_points = st.number_input("🔢 Number of Compaction Trials", 3, 10, st.session_state.lc_num_points)
    if num_points != st.session_state.lc_num_points:
        st.session_state.lc_num_points = num_points
        st.session_state.lc_trial_inputs = [{"W1":0.0,"W2":0.0,"W3":0.0,"W4":0.0,"W5":0.0} for _ in range(num_points)]

    # Mould dimensions
    st.markdown("### 📏 Mould Dimensions")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.lc_mould_dia = st.number_input("Diameter (cm)", value=st.session_state.lc_mould_dia)
    with col2:
        st.session_state.lc_mould_height = st.number_input("Height (cm)", value=st.session_state.lc_mould_height)

    volume = (math.pi/4)*st.session_state.lc_mould_dia**2*st.session_state.lc_mould_height
    st.info(f"📐 Volume of Mould = {volume:.2f} cm³")

    # Trial Inputs
    st.markdown("### 📋 Enter Trial Data")
    for i in range(num_points):
        st.markdown(f"#### Trial {i+1}")
        c1, c2 = st.columns(2)

        with c1:
            st.session_state.lc_trial_inputs[i]["W1"] = st.number_input(f"W1 (g)", value=st.session_state.lc_trial_inputs[i]["W1"], key=f"W1{i}")
            st.session_state.lc_trial_inputs[i]["W2"] = st.number_input(f"W2 (g)", value=st.session_state.lc_trial_inputs[i]["W2"], key=f"W2{i}")
            st.session_state.lc_trial_inputs[i]["W3"] = st.number_input(f"W3 (g)", value=st.session_state.lc_trial_inputs[i]["W3"], key=f"W3{i}")

        with c2:
            st.session_state.lc_trial_inputs[i]["W4"] = st.number_input(f"W4 (g)", value=st.session_state.lc_trial_inputs[i]["W4"], key=f"W4{i}")
            st.session_state.lc_trial_inputs[i]["W5"] = st.number_input(f"W5 (g)", value=st.session_state.lc_trial_inputs[i]["W5"], key=f"W5{i}")

    # RESET BUTTON
    if st.button("🔄 Reset Inputs"):
        st.session_state.lc_trial_inputs = [{"W1":0.0,"W2":0.0,"W3":0.0,"W4":0.0,"W5":0.0} for _ in range(num_points)]

    # CALCULATE
    if st.button("Calculate Compaction Results"):
        calculated_data = []

        for i, t in enumerate(st.session_state.lc_trial_inputs):
            W1, W2, W3, W4, W5 = t["W1"], t["W2"], t["W3"], t["W4"], t["W5"]

            if W2 > W1 and W3 > W1 and W5 > W4:
                weight_wet = W2 - W1
                weight_dry = W3 - W1

                if weight_dry == 0:
                    st.warning(f"Trial {i+1}: Dry weight cannot be zero")
                    continue

                water_content = ((weight_wet - weight_dry)/weight_dry)*100

                if water_content < 0 or water_content > 100:
                    st.warning(f"Trial {i+1}: Unusual water content")

                wet_density = (W5 - W4)/volume
                dry_density = wet_density/(1 + water_content/100)

                calculated_data.append({
                    "Trial": i+1,
                    "Water Content (%)": round(water_content,2),
                    "Dry Density (g/cc)": round(dry_density,3)
                })

        if calculated_data:
            df = pd.DataFrame(calculated_data).sort_values(by="Water Content (%)")

            mdd = df["Dry Density (g/cc)"].max()
            omc = df.loc[df["Dry Density (g/cc)"].idxmax(),"Water Content (%)"]

            # Plot
            fig, ax = plt.subplots()
            ax.plot(df["Water Content (%)"], df["Dry Density (g/cc)"], marker='o')
            ax.set_xlabel("Water Content (%)")
            ax.set_ylabel("Dry Density (g/cc)")
            ax.set_title("Compaction Curve")
            ax.grid(True)

            ax.plot(omc, mdd, 'ro')
            ax.axvline(omc, linestyle='--')
            ax.axhline(mdd, linestyle='--')

            img_buf = BytesIO()
            fig.savefig(img_buf, format="png", bbox_inches="tight")
            img_buf.seek(0)
            plt.close(fig)

            # SAVE RESULTS
            st.session_state.lc_results = {
                "df": df,
                "mdd": mdd,
                "omc": omc,
                "img_buf": img_buf,
                "volume": volume
            }

    # DISPLAY RESULTS (PERSISTENT)
    if "lc_results" in st.session_state:
        results = st.session_state.lc_results
        df = results["df"]
        mdd = results["mdd"]
        omc = results["omc"]

        st.markdown("### 📊 Results")
        st.dataframe(df)

        st.success(f"Maximum Dry Density (MDD) = {mdd:.3f} g/cc")
        st.success(f"Optimum Moisture Content (OMC) = {omc:.2f}%")

        st.image(results["img_buf"])

        # WORD REPORT
        if st.button("📄 Generate Word Report"):
            doc = Document()
            doc.add_heading("Light Compaction Test Report", 0)

            doc.add_paragraph(f"Mould Volume: {results['volume']:.2f} cm³")
            doc.add_paragraph(f"MDD: {mdd:.3f} g/cc")
            doc.add_paragraph(f"OMC: {omc:.2f}%")

            doc.add_heading("Compaction Curve", 1)
            doc.add_picture(results["img_buf"], width=Inches(5))

            doc.add_heading("Results Table", 1)
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'

            for i, col in enumerate(df.columns):
                table.rows[0].cells[i].text = col

            for i in range(len(df)):
                row = table.add_row().cells
                for j, col in enumerate(df.columns):
                    row[j].text = str(df[col].iloc[i])

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download Word Report",
                data=buffer,
                file_name="Light_Compaction_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    run()