import streamlit as st
import numpy as np
import pandas as pd
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches

def run():
    st.subheader("Specific Gravity by Density Bottle Method (IS 2720: Part 3, Sec 1: 1980)")

    # -----------------------------
    # Procedure
    # -----------------------------
    with st.expander("📝 Procedure & Formula", expanded=True):
        st.markdown("""
### Test Procedure:
1. Weigh the empty density bottle (W1).
2. Add dry soil and weigh the bottle (W2).
3. Fill the bottle with carbon tetrachloride (CCl₄) and weigh (W3).
4. Weigh the bottle with CCl₄ only (W4).
5. Calculate specific gravity (G).

### Formula:
G = ((W2 - W1) × ((W4 - W1) / V)) / ((W4 - W1) - (W3 - W2))
""")

    # -----------------------------
    # Session State Initialization
    # -----------------------------
    if "sg_volume" not in st.session_state:
        st.session_state.sg_volume = 50.0

    if "sg_num_trials" not in st.session_state:
        st.session_state.sg_num_trials = 3

    if "sg_trial_inputs" not in st.session_state:
        st.session_state.sg_trial_inputs = [{"W1":0.0,"W2":0.0,"W3":0.0,"W4":0.0} for _ in range(3)]

    if "sg_results" not in st.session_state:
        st.session_state.sg_results = None

    if "sg_avg" not in st.session_state:
        st.session_state.sg_avg = None

    if "sg_soil_type" not in st.session_state:
        st.session_state.sg_soil_type = None

    # -----------------------------
    # Bottle Volume
    # -----------------------------
    st.session_state.sg_volume = st.number_input(
        "Volume of Density Bottle (cm³)",
        min_value=0.1,
        value=float(st.session_state.sg_volume),
        step=0.1,
        format="%.3f"
    )

    # -----------------------------
    # Number of Trials
    # -----------------------------
    num_trials = st.number_input(
        "Number of Trials",
        min_value=1,
        max_value=10,
        value=st.session_state.sg_num_trials,
        step=1
    )

    if num_trials != st.session_state.sg_num_trials:
        st.session_state.sg_num_trials = num_trials
        st.session_state.sg_trial_inputs = [{"W1":0.0,"W2":0.0,"W3":0.0,"W4":0.0} for _ in range(num_trials)]

    # -----------------------------
    # Input Fields
    # -----------------------------
    st.markdown("### Enter Data for Each Trial")

    for i in range(num_trials):
        st.markdown(f"#### Trial {i+1}")
        col1, col2 = st.columns(2)

        with col1:
            st.session_state.sg_trial_inputs[i]["W1"] = st.number_input(
                f"W1 (Empty) [{i+1}]",
                value=float(st.session_state.sg_trial_inputs[i]["W1"]),
                step=0.01,
                format="%.3f",
                key=f"sg_W1_{i}"
            )

            st.session_state.sg_trial_inputs[i]["W2"] = st.number_input(
                f"W2 (Bottle + Dry Soil) [{i+1}]",
                value=float(st.session_state.sg_trial_inputs[i]["W2"]),
                step=0.01,
                format="%.3f",
                key=f"sg_W2_{i}"
            )

        with col2:
            st.session_state.sg_trial_inputs[i]["W3"] = st.number_input(
                f"W3 (Bottle + Soil + CCl₄) [{i+1}]",
                value=float(st.session_state.sg_trial_inputs[i]["W3"]),
                step=0.01,
                format="%.3f",
                key=f"sg_W3_{i}"
            )

            st.session_state.sg_trial_inputs[i]["W4"] = st.number_input(
                f"W4 (Bottle + CCl₄) [{i+1}]",
                value=float(st.session_state.sg_trial_inputs[i]["W4"]),
                step=0.01,
                format="%.3f",
                key=f"sg_W4_{i}"
            )

    # -----------------------------
    # Save Inputs CSV
    # -----------------------------
    if st.button("💾 Save Inputs"):
        df_save = pd.DataFrame(st.session_state.sg_trial_inputs)
        df_save.insert(0, "Trial", range(1,len(df_save)+1))
        buffer = StringIO()
        df_save.to_csv(buffer,index=False)
        buffer.seek(0)

        st.download_button(
            "📥 Download Input Data as CSV",
            data=buffer.getvalue(),
            file_name="specific_gravity_inputs.csv",
            mime="text/csv"
        )

    # -----------------------------
    # Calculate
    # -----------------------------
    if st.button("Calculate Specific Gravity"):
        results = []
        valid_G_list = []
        V = st.session_state.sg_volume

        for i, trial in enumerate(st.session_state.sg_trial_inputs):
            W1, W2, W3, W4 = trial["W1"], trial["W2"], trial["W3"], trial["W4"]

            if not (W2>W1 and W4>W1 and W3>=W2):
                results.append({"Trial": i+1, "Gc": None, "G": None})
                continue

            Gc = (W4 - W1) / V
            denom = (W4 - W1) - (W3 - W2)

            if denom != 0:
                G = (W2 - W1) * Gc / denom
                valid_G_list.append(G)
                results.append({"Trial": i+1, "Gc": round(Gc,3), "G": round(G,3)})
            else:
                results.append({"Trial": i+1, "Gc": round(Gc,3), "G": None})

        if valid_G_list:
            df_results = pd.DataFrame(results)
            G_avg = sum(valid_G_list)/len(valid_G_list)

            st.session_state.sg_results = df_results
            st.session_state.sg_avg = G_avg

            # Soil interpretation
            if G_avg < 2.60:
                soil_type = "Organic soil"
            elif 2.60 <= G_avg <= 2.67:
                soil_type = "Sand / inorganic soil"
            elif 2.67 < G_avg <= 2.78:
                soil_type = "Silty sand / clay"
            else:
                soil_type = "Dense clay / heavy minerals"

            st.session_state.sg_soil_type = soil_type

    # -----------------------------
    # Display Stored Results
    # -----------------------------
    if st.session_state.sg_results is not None:

        st.markdown("### Results")
        st.dataframe(st.session_state.sg_results)

        st.success(f"Average Specific Gravity = {st.session_state.sg_avg:.3f}")
        st.info(f"Soil Type: {st.session_state.sg_soil_type}")

        # -----------------------------
        # Word Report
        # -----------------------------
        if st.button("📄 Generate Word Report"):

            doc = Document()
            doc.add_heading("Specific Gravity Test Report",0)

            doc.add_paragraph(f"Bottle Volume: {st.session_state.sg_volume} cm³")
            doc.add_paragraph(f"Average Specific Gravity: {st.session_state.sg_avg:.3f}")
            doc.add_paragraph(f"Soil Type: {st.session_state.sg_soil_type}")

            doc.add_heading("Trial Results", level=1)

            df_results = st.session_state.sg_results
            table = doc.add_table(rows=1, cols=len(df_results.columns))

            for i, col in enumerate(df_results.columns):
                table.rows[0].cells[i].text = col

            for i in range(len(df_results)):
                row_cells = table.add_row().cells
                for j, col in enumerate(df_results.columns):
                    row_cells[j].text = str(df_results[col].iloc[i])

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download Report",
                data=buffer,
                file_name="Specific_Gravity_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    return None