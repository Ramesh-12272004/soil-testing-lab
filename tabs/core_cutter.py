import streamlit as st
import numpy as np
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime

def run():
    st.subheader("🧪 In-situ Density by Core Cutter Method")
    st.caption("IS 2720: Part 29: 1975, reaffirmed 1995")

    # ---------------- PROCEDURE ----------------
    with st.expander("📘 View Detailed Procedure"):
        st.markdown("""
### Objective
To determine the bulk and dry density of soil in-situ using the core cutter method.

### Formulae
- Bulk Density: ρb = W / V
- Moisture Content: w = [(Ww - Wd) / Wd] × 100
- Dry Density: ρd = ρb / (1 + w/100)
""")

    # ---------------- INPUTS ----------------
    st.markdown("### 📏 Core Cutter Dimensions")
    h = st.number_input("Height of Core Cutter (cm)", value=12.8, min_value=0.1, key="cc_h")
    d = st.number_input("Internal Diameter (cm)", value=10.0, min_value=0.1, key="cc_d")

    volume = np.pi * (d / 2) ** 2 * h
    st.info(f"Volume = {volume:.2f} cm³")

    st.markdown("### ⚖️ Weight Measurements")

    wt_empty = st.number_input("Empty Cutter (g)", value=0.0, key="cc_wt_empty")
    wt_full = st.number_input("Cutter + Soil (g)", value=max(wt_empty,0.0), min_value=wt_empty, key="cc_wt_full")
    wt_container = st.number_input("Empty Container (g)", value=0.0, key="cc_wt_container")
    wt_wet = st.number_input("Container + Wet Soil (g)", value=max(wt_container,0.0), min_value=wt_container, key="cc_wt_wet")
    wt_dry = st.number_input("Container + Dry Soil (g)", value=max(wt_container,0.0), min_value=wt_container, key="cc_wt_dry")

    # ---------------- CALCULATION ----------------
    if st.button("📊 Calculate Core Cutter Results"):

        if not (h > 0 and d > 0 and wt_full > wt_empty and wt_wet > wt_container and wt_dry > wt_container):
            st.error("Enter valid inputs")
            return

        wt_soil = wt_full - wt_empty
        bulk_density = wt_soil / volume

        Wd = wt_dry - wt_container
        Ww = wt_wet - wt_container
        moisture_content = ((Ww - Wd) / Wd) * 100 if Wd > 0 else 0

        dry_density = bulk_density / (1 + moisture_content / 100)

        if dry_density < 1.4:
            suitability = "Low compaction"
        elif dry_density < 1.75:
            suitability = "Moderate compaction"
        else:
            suitability = "Well compacted"

        # ✅ STORE EVERYTHING
        st.session_state.cc_results = {
            "h": h, "d": d, "volume": volume,
            "wt_empty": wt_empty,
            "wt_full": wt_full,
            "wt_container": wt_container,
            "wt_wet": wt_wet,
            "wt_dry": wt_dry,
            "bulk_density": bulk_density,
            "moisture_content": moisture_content,
            "dry_density": dry_density,
            "suitability": suitability
        }

    # ---------------- DISPLAY RESULTS ----------------
    if "cc_results" in st.session_state:
        res = st.session_state.cc_results

        st.markdown("### Results")
        st.info(f"Bulk Density = {res['bulk_density']:.2f}")
        st.info(f"Moisture Content = {res['moisture_content']:.2f}%")
        st.success(f"Dry Density = {res['dry_density']:.2f}")

        st.markdown("### Suitability")
        st.write(res["suitability"])

        # ---------------- WORD REPORT ----------------
        if st.button("📄 Generate Word Report"):

            doc = Document()
            doc.add_heading("Core Cutter Test Report", 0)

            doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

            # INPUT DATA TABLE
            doc.add_heading("Input Data", 1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            inputs = [
                ("Height (cm)", res["h"]),
                ("Diameter (cm)", res["d"]),
                ("Volume (cm³)", round(res["volume"],2)),
                ("Empty Cutter (g)", res["wt_empty"]),
                ("Cutter + Soil (g)", res["wt_full"]),
                ("Empty Container (g)", res["wt_container"]),
                ("Wet Soil (g)", res["wt_wet"]),
                ("Dry Soil (g)", res["wt_dry"])
            ]

            table.rows[0].cells[0].text = "Parameter"
            table.rows[0].cells[1].text = "Value"

            for name, val in inputs:
                row = table.add_row().cells
                row[0].text = name
                row[1].text = str(val)

            # RESULTS TABLE
            doc.add_heading("Results", 1)
            table2 = doc.add_table(rows=1, cols=2)
            table2.style = 'Table Grid'

            results = [
                ("Bulk Density", round(res["bulk_density"],2)),
                ("Moisture Content (%)", round(res["moisture_content"],2)),
                ("Dry Density", round(res["dry_density"],2))
            ]

            table2.rows[0].cells[0].text = "Parameter"
            table2.rows[0].cells[1].text = "Value"

            for name, val in results:
                row = table2.add_row().cells
                row[0].text = name
                row[1].text = str(val)

            # REMARKS
            doc.add_heading("Remarks", 1)
            doc.add_paragraph(res["suitability"])

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download Word Report",
                data=buffer.getvalue(),
                file_name="Core_Cutter_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    return None