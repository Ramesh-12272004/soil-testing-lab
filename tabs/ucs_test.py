import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import math

def generate_word_report(df, mean_ucs):
    doc = Document()
    doc.add_heading('Unconfined Compressive Strength (UCS) Test Report', 0)
    doc.add_paragraph('IS Code: IS 2720 Part 10:1991')

    # Add procedure
    doc.add_heading('Procedure', level=1)
    doc.add_paragraph(
        "1. Prepare cylindrical soil specimens with uniform diameter and length.\n"
        "2. Place the specimen in the testing machine.\n"
        "3. Apply axial load using the proving ring until failure.\n"
        "4. Record the proving ring reading at failure.\n"
        "5. Repeat for all trials."
    )

    # Add formulas
    doc.add_heading('Formulas Used', level=1)
    doc.add_paragraph(
        "Cross-sectional Area: A = π × (Diameter / 2)²\n"
        "Failure Load: P = Proving Ring Constant × Ring Reading\n"
        "Unconfined Compressive Strength (UCS): UCS = P / A × 1000 (kPa)"
    )

    doc.add_paragraph(f"Average UCS: {round(mean_ucs, 2)} kPa")

    # Add results table
    doc.add_heading('Results Table', level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Add plot image
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
        fig, ax = plt.subplots()
        ax.bar(df["Trial"], df["UCS (kPa)"], color="lightgreen")
        ax.set_xlabel("Trial")
        ax.set_ylabel("UCS (kPa)")
        ax.set_title("Unconfined Compressive Strength")
        plt.tight_layout()
        fig.savefig(tmpfile.name)
        doc.add_paragraph("\nUCS Chart:")
        doc.add_picture(tmpfile.name, width=Inches(5))
        plt.close(fig)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def run():
    st.subheader("🧪 Unconfined Compressive Strength (UCS) Test - IS 2720 Part 10:1991")

    st.markdown("""
    ### 🔹 Purpose
    To determine the **Unconfined Compressive Strength (UCS)** of cohesive soils.
    
    ### 🔹 Procedure
    1. Prepare cylindrical soil specimens of known diameter and length.  
    2. Place the specimen in the UCS testing machine.  
    3. Apply axial load using a proving ring until the specimen fails.  
    4. Record the proving ring reading at failure.  
    5. Repeat for all trials.
    
    ### 🔹 Formulas Used
    - Cross-sectional Area: A = π × (Diameter / 2)²  
    - Failure Load: P = Proving Ring Constant × Ring Reading  
    - Unconfined Compressive Strength: UCS = P / A × 1000 (kPa)
    """)

    # --- Session State Initialization ---
    if "ucs_num_trials" not in st.session_state:
        st.session_state.ucs_num_trials = 3
    if "ucs_inputs" not in st.session_state:
        st.session_state.ucs_inputs = []

    # Number of Trials
    num_trials = st.number_input(
        "Number of Trials",
        min_value=1, max_value=10,
        value=st.session_state.ucs_num_trials,
        step=1
    )
    if num_trials != st.session_state.ucs_num_trials:
        st.session_state.ucs_num_trials = num_trials
        st.session_state.ucs_inputs = [{"d":0.0,"l":0.0,"k":0.0,"r":0.0} for _ in range(num_trials)]

    # Initialize inputs if empty
    while len(st.session_state.ucs_inputs) < num_trials:
        st.session_state.ucs_inputs.append({"d":0.0,"l":0.0,"k":0.0,"r":0.0})
    st.session_state.ucs_inputs = st.session_state.ucs_inputs[:num_trials]

    # Input for each trial
    results = []
    for i in range(num_trials):
        st.markdown(f"### Trial {i+1}")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.session_state.ucs_inputs[i]["d"] = st.number_input(f"Diameter (mm)", key=f"d_{i}", min_value=0.0, value=st.session_state.ucs_inputs[i]["d"])
        with col2:
            st.session_state.ucs_inputs[i]["l"] = st.number_input(f"Length (mm)", key=f"l_{i}", min_value=0.0, value=st.session_state.ucs_inputs[i]["l"])
        with col3:
            st.session_state.ucs_inputs[i]["k"] = st.number_input(f"Proving Ring Constant (N/div)", key=f"k_{i}", min_value=0.0, value=st.session_state.ucs_inputs[i]["k"])
        with col4:
            st.session_state.ucs_inputs[i]["r"] = st.number_input(f"Ring Reading (div)", key=f"r_{i}", min_value=0.0, value=st.session_state.ucs_inputs[i]["r"])

        # Calculate UCS for valid inputs
        d, k, r = st.session_state.ucs_inputs[i]["d"], st.session_state.ucs_inputs[i]["k"], st.session_state.ucs_inputs[i]["r"]
        if d>0 and k>0 and r>0:
            area = math.pi*(d/2)**2
            failure_load = k * r
            ucs = (failure_load / area)*1000
            results.append({
                "Trial": i+1,
                "Diameter (mm)": d,
                "Length (mm)": st.session_state.ucs_inputs[i]["l"],
                "Area (mm²)": round(area,2),
                "Proving Ring Constant (N/div)": k,
                "Ring Reading (div)": r,
                "Failure Load (N)": round(failure_load,2),
                "UCS (kPa)": round(ucs,2)
            })

    # --- Calculate Button ---
    if st.button("📊 Calculate UCS"):
        if results:
            df = pd.DataFrame(results)
            st.markdown("### 📋 Result Table")
            st.dataframe(df)

            mean_ucs = df["UCS (kPa)"].mean()
            st.success(f"🔹 Average UCS = {round(mean_ucs, 2)} kPa")

            st.markdown("### 📈 UCS per Trial")
            fig, ax = plt.subplots()
            ax.bar(df["Trial"], df["UCS (kPa)"], color="lightgreen")
            ax.set_xlabel("Trial")
            ax.set_ylabel("UCS (kPa)")
            ax.set_title("Unconfined Compressive Strength")
            st.pyplot(fig)
            plt.close(fig)

            # Download Word report
            buffer = generate_word_report(df, mean_ucs)
            st.download_button(
                label="📄 Download UCS Word Report",
                data=buffer,
                file_name="UCS_Test_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Please enter valid values for all required fields.")