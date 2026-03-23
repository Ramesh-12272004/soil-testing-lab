import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

def run():
    st.subheader("🧪 Undrained Triaxial Test (IS 2720 Part 11:1971)")

    # --- Procedure & Formula ---
    with st.expander("📝 Procedure & Formula", expanded=True):
        st.markdown("""
### Test Procedure:
1. Prepare cylindrical soil sample.
2. Apply confining pressure σ₃.
3. Apply axial load until failure and record deviator stress (σ₁ - σ₃).
4. Measure failure angle if required.
5. Calculate major and minor principal stresses:
   
**Principal Stresses (Text Formula):**  
- σ₁ = Confining Pressure + Deviator Stress  
- σ₃ = Confining Pressure  

6. Determine cohesion (c) and angle of internal friction (ϕ) from Mohr’s circles and failure envelope.
""")

    st.markdown("### 🧾 Enter Sample Details")
    sample_area = st.number_input("Cross-sectional Area of Sample (cm²)", value=38.0, min_value=1.0)
    sample_length = st.number_input("Length of Sample (cm)", value=7.6, min_value=1.0)

    st.markdown("### 🔢 Enter Number of Trials")
    num_trials = st.number_input("Number of Trials", min_value=1, max_value=10, value=3, step=1)

    st.markdown("### ✍️ Enter Data for Each Trial")
    trial_data = []
    for i in range(num_trials):
        st.markdown(f"#### Trial {i+1}")
        col1, col2, col3 = st.columns(3)
        with col1:
            confining_pressure = st.number_input(f"Confining Pressure (σ₃) [kg/cm²] [{i+1}]", key=f"cp_{i}", min_value=0.0)
        with col2:
            deviator_stress = st.number_input(f"Deviator Stress (σ₁ - σ₃) [kg/cm²] [{i+1}]", key=f"ds_{i}", min_value=0.0)
        with col3:
            failure_angle = st.number_input(f"Failure Angle (°) [{i+1}]", value=90.0, key=f"fa_{i}")
        trial_data.append((confining_pressure, deviator_stress, failure_angle))

    if st.button("📊 Calculate Results"):
        df = pd.DataFrame(trial_data, columns=["Confining Pressure", "Deviator Stress", "Failure Angle"])
        df["Major Principal Stress (σ₁)"] = df["Confining Pressure"] + df["Deviator Stress"]
        df["Minor Principal Stress (σ₃)"] = df["Confining Pressure"]

        # Calculate Mohr circles
        sigma1 = df["Major Principal Stress (σ₁)"]
        sigma3 = df["Minor Principal Stress (σ₃)"]
        center = (sigma1 + sigma3) / 2
        radius = (sigma1 - sigma3) / 2

        # Plot Mohr circles
        st.markdown("### 📈 Mohr’s Circles")
        fig, ax = plt.subplots(figsize=(6,6))
        for c, r in zip(center, radius):
            circle = plt.Circle((c, 0), r, fill=False)
            ax.add_patch(circle)
            ax.plot(c + r, 0, 'ro')  # σ₁
            ax.plot(c - r, 0, 'bo')  # σ₃

        ax.set_xlabel("Normal Stress (σ) [kg/cm²]")
        ax.set_ylabel("Shear Stress (τ) [kg/cm²]")
        ax.set_title("Mohr’s Circles for Triaxial Test")
        ax.axis('equal')
        ax.grid(True)
        ax.set_xlim(left=0)
        st.pyplot(fig)

        # Fit failure envelope
        normal_stress = center
        shear_stress = radius  # Approximation for linear envelope
        fit = np.polyfit(normal_stress, shear_stress, 1)
        cohesion = fit[1]
        phi_deg = np.degrees(np.arctan(fit[0]))

        st.success(f"**Cohesion (c)** ≈ {cohesion:.4f} kg/cm²")
        st.success(f"**Angle of Internal Friction (ϕ)** ≈ {phi_deg:.2f}°")

        st.markdown("### 📋 Trial Results Table")
        st.dataframe(df.style.format(precision=3))

        # --- Word Report Generation ---
        if st.button("📥 Download Word Report"):
            doc = Document()
            doc.add_heading("Undrained Triaxial Test Report", 0)
            doc.add_paragraph(f"Sample Area: {sample_area} cm²")
            doc.add_paragraph(f"Sample Length: {sample_length} cm")
            doc.add_paragraph(f"Cohesion (c): {cohesion:.4f} kg/cm²")
            doc.add_paragraph(f"Angle of Internal Friction (ϕ): {phi_deg:.2f}°")

            doc.add_heading("Trial Data", level=1)
            table = doc.add_table(rows=1, cols=len(df.columns)+1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Trial"
            for idx, col in enumerate(df.columns):
                hdr_cells[idx+1].text = col

            for i, row in df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(i+1)
                for j, col in enumerate(df.columns):
                    row_cells[j+1].text = f"{row[col]:.3f}"

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="📥 Download Triaxial Report",
                data=buffer,
                file_name="Undrained_Triaxial_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )