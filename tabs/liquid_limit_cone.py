import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches

def run():
    st.subheader("Liquid Limit by Cone Penetration Method (IS 2720: Part 5: 1985)")

    # --- Procedure ---
    with st.expander("📝 Procedure & Formulas", expanded=True):
        st.markdown("""
W1 = Empty container  
W2 = Container + wet soil  
W3 = Container + dry soil  

Water Content = ((W2 - W3)/(W3 - W1)) × 100  
Liquid Limit = Water content at 20 mm penetration  
""")

    def calculate_moisture_content(w1, w2, w3):
        if not (w2 > w3 > w1): return np.nan
        return ((w2 - w3) / (w3 - w1)) * 100

    # ---------------- INPUT ----------------
    if "cone_trial_inputs" not in st.session_state:
        st.session_state.cone_trial_inputs = [
            {"penetration":0.0,"w1":0.0,"w2":0.0,"w3":0.0,"water_content":0.0}
            for _ in range(4)
        ]

    num_trials = st.number_input("Trials", 3, 10, 4)

    for i in range(num_trials):
        st.markdown(f"### Trial {i+1}")
        data = st.session_state.cone_trial_inputs[i]

        data["penetration"] = st.number_input(f"Penetration {i}", value=data["penetration"])
        data["w1"] = st.number_input(f"W1 {i}", value=data["w1"])
        data["w2"] = st.number_input(f"W2 {i}", value=data["w2"])
        data["w3"] = st.number_input(f"W3 {i}", value=data["w3"])

        mc = calculate_moisture_content(data["w1"], data["w2"], data["w3"])
        data["water_content"] = mc

        if np.isnan(mc):
            st.error("Invalid input")
        else:
            st.info(f"Water Content = {mc:.2f}%")

    # ---------------- CALCULATE ----------------
    if st.button("Calculate Liquid Limit"):

        df_all = pd.DataFrame(st.session_state.cone_trial_inputs)
        df_all.insert(0,"Trial", range(1,len(df_all)+1))

        df_valid = df_all[(df_all["penetration"]>0) & (~df_all["water_content"].isna())]

        if len(df_valid) < 2:
            st.error("Need at least 2 valid points")
            return

        try:
            x = df_valid["penetration"]
            y = df_valid["water_content"]

            coeffs = np.polyfit(x,y,1)
            poly = np.poly1d(coeffs)

            liquid_limit = poly(20)

            fig, ax = plt.subplots()
            ax.scatter(x,y)
            xp = np.linspace(min(x), max(x),100)
            ax.plot(xp, poly(xp))
            ax.plot(20, liquid_limit, 'ro')

            st.pyplot(fig)

            img_buf = BytesIO()
            fig.savefig(img_buf)
            img_buf.seek(0)

            # ✅ STORE RESULTS
            st.session_state.cone_results = {
                "df_all": df_all,
                "liquid_limit": liquid_limit,
                "img_buf": img_buf,
                "soil_class": "Low" if liquid_limit<35 else "Medium" if liquid_limit<=50 else "High"
            }

        except np.linalg.LinAlgError:
            st.error("Curve fit error")
        except Exception as e:
            st.error(str(e))

    # ---------------- RESULT DISPLAY ----------------
    if "cone_results" in st.session_state:

        res = st.session_state.cone_results

        st.success(f"Liquid Limit = {res['liquid_limit']:.2f}%")
        st.image(res["img_buf"])

        # ---------------- WORD REPORT ----------------
        if st.button("📄 Generate Word Report"):

            doc = Document()
            doc.add_heading("Liquid Limit Report (Cone Method)",0)

            # Procedure
            doc.add_heading("Procedure",1)
            doc.add_paragraph("""
W1 = Empty container  
W2 = Wet soil  
W3 = Dry soil  
LL determined at 20 mm penetration
""")

            # Input Data
            doc.add_heading("Input Data",1)
            df_all = res["df_all"]

            table = doc.add_table(rows=1, cols=len(df_all.columns))
            for i,col in enumerate(df_all.columns):
                table.rows[0].cells[i].text = col

            for i in range(len(df_all)):
                row = table.add_row().cells
                for j,col in enumerate(df_all.columns):
                    row[j].text = str(df_all[col].iloc[i])

            # Results
            doc.add_heading("Results",1)
            doc.add_paragraph(f"Liquid Limit = {res['liquid_limit']:.2f}%")
            doc.add_paragraph(f"Soil Type = {res['soil_class']}")

            # Graph
            doc.add_picture(res["img_buf"], width=Inches(6))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "Download Report",
                buffer.getvalue(),
                "cone_LL_report.docx"
            )

    return None