import streamlit as st
import math
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches

def run():
    st.subheader("Variable Head Permeability Test (IS 2720 Part 17:1986)")

    # ---------------- SESSION STATE ----------------
    if "vh_results" not in st.session_state:
        st.session_state.vh_results = None
    if "vh_avg_k" not in st.session_state:
        st.session_state.vh_avg_k = None
    if "vh_soil" not in st.session_state:
        st.session_state.vh_soil = None
    if "vh_graph" not in st.session_state:
        st.session_state.vh_graph = None

    # ---------------- PROCEDURE ----------------
    with st.expander("📝 Procedure & Formula", expanded=True):
        st.markdown("""
k = (2.3 × a × L / (A × t)) × log₁₀(h₁ / h₂)
""")

    # ---------------- INITIAL STATE ----------------
    if "vh_num_trials" not in st.session_state:
        st.session_state.vh_num_trials = 3
    if "vh_a" not in st.session_state:
        st.session_state.vh_a = 1.0
    if "vh_A" not in st.session_state:
        st.session_state.vh_A = 50.0
    if "vh_L" not in st.session_state:
        st.session_state.vh_L = 10.0

    # ---------------- TRIAL COUNT ----------------
    num_trials = st.number_input("Number of Trials", 1, 10, st.session_state.vh_num_trials)

    if num_trials != st.session_state.vh_num_trials:
        st.session_state.vh_num_trials = num_trials
        st.session_state.vh_trial_inputs = [{"h1":0.0,"h2":0.0,"t":0.0} for _ in range(num_trials)]

    if "vh_trial_inputs" not in st.session_state:
        st.session_state.vh_trial_inputs = [{"h1":0.0,"h2":0.0,"t":0.0} for _ in range(num_trials)]

    # ---------------- CONSTANTS ----------------
    with st.expander("Enter Constant Parameters"):
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.vh_a = st.number_input("a (cm²)", value=float(st.session_state.vh_a), step=0.01)
            st.session_state.vh_A = st.number_input("A (cm²)", value=float(st.session_state.vh_A), step=0.01)
        with col2:
            st.session_state.vh_L = st.number_input("L (cm)", value=float(st.session_state.vh_L), step=0.01)

    # ---------------- INPUT ----------------
    for i in range(num_trials):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.session_state.vh_trial_inputs[i]["h1"] = st.number_input(
                f"h1 - T{i+1}", value=float(st.session_state.vh_trial_inputs[i]["h1"]), step=0.01
            )
        with col2:
            st.session_state.vh_trial_inputs[i]["h2"] = st.number_input(
                f"h2 - T{i+1}", value=float(st.session_state.vh_trial_inputs[i]["h2"]), step=0.01
            )
        with col3:
            st.session_state.vh_trial_inputs[i]["t"] = st.number_input(
                f"t - T{i+1}", value=float(st.session_state.vh_trial_inputs[i]["t"]), step=0.01
            )

    # ---------------- SAVE INPUT ----------------
    if st.button("💾 Save Inputs"):
        df = pd.DataFrame(st.session_state.vh_trial_inputs)
        buffer = StringIO()
        df.to_csv(buffer, index=False)
        buffer.seek(0)
        st.download_button("Download CSV", buffer.getvalue())

    # ---------------- CALCULATE ----------------
    if st.button("Calculate Permeability"):

        a = st.session_state.vh_a
        A = st.session_state.vh_A
        L = st.session_state.vh_L

        results = []

        for i, trial in enumerate(st.session_state.vh_trial_inputs):
            h1, h2, t = trial["h1"], trial["h2"], trial["t"]

            if h1 > h2 > 0 and t > 0:
                k = (2.3 * a * L) / (A * t) * math.log10(h1 / h2)
                results.append({
                    "Trial": i+1,
                    "h1": h1,
                    "h2": h2,
                    "t": t,
                    "k": round(k,6)
                })

        if results:
            df = pd.DataFrame(results)
            avg_k = df["k"].mean()

            # Soil classification
            if avg_k < 1e-7:
                soil = "Clay"
            elif avg_k < 1e-5:
                soil = "Silty Clay"
            elif avg_k < 1e-3:
                soil = "Silty Sand"
            else:
                soil = "Sand/Gravel"

            # Graph
            fig, ax = plt.subplots()
            ax.plot(df["t"], df["k"], marker='o')
            ax.set_xlabel("Time")
            ax.set_ylabel("k")
            ax.grid(True)

            img = BytesIO()
            fig.savefig(img, format="png")
            img.seek(0)

            # SAVE STATE
            st.session_state.vh_results = df
            st.session_state.vh_avg_k = avg_k
            st.session_state.vh_soil = soil
            st.session_state.vh_graph = img

    # ---------------- DISPLAY ----------------
    if st.session_state.vh_results is not None:

        st.dataframe(st.session_state.vh_results)
        st.success(f"Average k = {st.session_state.vh_avg_k:.6f}")
        st.info(f"Soil: {st.session_state.vh_soil}")

        st.image(st.session_state.vh_graph)

        # ---------------- REPORT ----------------
        if st.button("📄 Generate Word Report"):

            df = st.session_state.vh_results
            avg_k = st.session_state.vh_avg_k
            soil = st.session_state.vh_soil
            img = st.session_state.vh_graph

            doc = Document()
            doc.add_heading("Variable Head Permeability Test Report", 0)

            doc.add_paragraph(f"Average k = {avg_k:.6f}")
            doc.add_paragraph(f"Soil = {soil}")

            table = doc.add_table(rows=1, cols=len(df.columns))
            for i, col in enumerate(df.columns):
                table.rows[0].cells[i].text = col

            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

            img.seek(0)
            doc.add_picture(img, width=Inches(5))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download Report",
                buffer,
                file_name="VH_Report.docx"
            )

    return None