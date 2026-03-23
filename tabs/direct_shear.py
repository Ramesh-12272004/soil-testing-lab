import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import math

def run():
    st.subheader("🔗 Direct Shear Test (IS 2720 Part 13:1986)")

    # --- Procedure and Formulas ---
    st.markdown("""
    ### Procedure
    1. Take a soil sample and place it in the shear box.
    2. Apply a normal load (sigma_n) on the sample using the loading mechanism.
    3. Measure horizontal deformation and proving ring readings at regular intervals.
    4. Determine the shear force at failure for each trial.
    5. Calculate shear stress, plot Shear Stress vs Deformation, and determine maximum shear stress (tau_max).
    6. Repeat for different normal stresses.
    7. Plot Normal Stress sigma_n vs Maximum Shear Stress tau_max to determine Cohesion (c) and Angle of Internal Friction (phi).

    ### Key Formulas (Plain Text)
    - Shear Force: Fs = Proving Ring Reading (div) * Proving Ring Constant (kg/div)
    - Shear Stress: tau = Fs / A, where A = shear area (cm²)
    - Deformation: delta = Horizontal Deformation (div) * Dial Gauge Least Count (mm/div)
    - Mohr-Coulomb Failure Criterion: tau = c + sigma_n * tan(phi)
      where:
        c = cohesion (kg/cm²)
        phi = angle of internal friction (degrees)
        sigma_n = applied normal stress (kg/cm²)
        tau = shear stress at failure (kg/cm²)
    """)

    # --- Input Initialization ---
    box_dim = st.number_input("Side Length of Shear Box (mm)", min_value=0.0, value=0.0, step=0.1, format="%.2f")
    area = (box_dim / 10) ** 2  # convert mm² to cm²
    st.info(f"Calculated Shear Area: {area:.2f} cm²")

    proving_ring_const = st.number_input("Proving Ring Constant (kg/div)", value=0.0, step=0.01, format="%.2f")
    dial_lc = st.number_input("Dial Gauge Least Count (mm/div)", value=0.0, step=0.001, format="%.3f")

    num_trials = st.number_input("Number of Normal Stress Trials", min_value=1, max_value=5, value=1, step=1)
    n_readings_per_trial = st.number_input("Number of Readings per Trial", min_value=2, value=5, step=1)

    # --- Horizontal Deformation (common to all trials) ---
    st.markdown("### Horizontal Deformation Readings (divisions)")
    horizontal_def = []
    for i in range(n_readings_per_trial):
        horizontal_def.append(st.number_input(f"Reading {i+1}", min_value=0.0, value=0.0, step=0.1, format="%.2f"))

    # --- Trial Inputs ---
    trials_data = []
    for t in range(num_trials):
        st.markdown(f"### Trial {t+1}")
        sigma_n = st.number_input(f"Normal Stress σₙ (kg/cm²) for Trial {t+1}", min_value=0.0, value=0.0, step=0.01, format="%.2f")
        proving_ring_readings = []
        st.markdown(f"Proving Ring Readings for Trial {t+1} (divisions)")
        for i in range(n_readings_per_trial):
            proving_ring_readings.append(st.number_input(f"PR Reading {i+1} - Trial {t+1}", min_value=0.0, value=0.0, step=0.01, format="%.2f"))
        trials_data.append({
            "sigma_n": sigma_n,
            "pr_readings": proving_ring_readings
        })

    # --- Calculate Button ---
    if st.button("🧮 Calculate Shear Test Results"):
        normal_stresses = []
        shear_stresses = []
        all_dfs = []

        for t, trial in enumerate(trials_data):
            sigma_n = trial["sigma_n"]
            pr_readings = trial["pr_readings"]

            df = pd.DataFrame({
                "Horizontal Deformation (div)": horizontal_def,
                "Proving Ring Reading (div)": pr_readings
            })
            df["Shear Force (kg)"] = df["Proving Ring Reading (div)"] * proving_ring_const
            df["Shear Stress (kg/cm²)"] = df["Shear Force (kg)"] / area
            df["Deformation (mm)"] = df["Horizontal Deformation (div)"] * dial_lc

            st.markdown(f"#### Trial {t+1} Data")
            st.dataframe(df)

            fig, ax = plt.subplots()
            ax.plot(df["Deformation (mm)"], df["Shear Stress (kg/cm²)"], marker='o')
            ax.set_xlabel("Deformation (mm)")
            ax.set_ylabel("Shear Stress (kg/cm²)")
            ax.set_title(f"Trial {t+1}: Shear Stress vs Deformation")
            ax.grid(True)
            st.pyplot(fig)
            plt.close(fig)

            tau_max = df["Shear Stress (kg/cm²)"].max()
            normal_stresses.append(sigma_n)
            shear_stresses.append(tau_max)
            all_dfs.append(df)

        # --- Cohesion & Friction Angle ---
        if len(normal_stresses) >= 2:
            coeffs = np.polyfit(normal_stresses, shear_stresses, 1)
            phi_deg = math.degrees(math.atan(coeffs[0]))
            cohesion = coeffs[1]

            st.success(f"**Cohesion (c): {cohesion:.3f} kg/cm²**")
            st.success(f"**Angle of Internal Friction (ϕ): {phi_deg:.2f}°**")

            fig2, ax2 = plt.subplots()
            ax2.scatter(normal_stresses, shear_stresses, color='blue')
            x_line = np.linspace(0, max(normal_stresses)*1.1, 100)
            y_line = coeffs[0]*x_line + coeffs[1]
            ax2.plot(x_line, y_line, color='red', label=f"Mohr-Coulomb Fit")
            ax2.set_xlabel("Normal Stress σₙ (kg/cm²)")
            ax2.set_ylabel("Shear Stress τ (kg/cm²)")
            ax2.grid(True)
            ax2.legend()
            st.pyplot(fig2)
            plt.close(fig2)

            # --- Word Report ---
            if st.button("📄 Generate Word Report"):
                doc = Document()
                doc.add_heading("Direct Shear Test Report", 0)
                doc.add_paragraph(f"Shear Box Side: {box_dim} mm")
                doc.add_paragraph(f"Proving Ring Constant: {proving_ring_const} kg/div")
                doc.add_paragraph(f"Dial Gauge LC: {dial_lc} mm/div")
                doc.add_paragraph(f"Shear Area: {area:.2f} cm²")
                doc.add_paragraph(f"Cohesion (c): {cohesion:.3f} kg/cm²")
                doc.add_paragraph(f"Angle of Internal Friction (ϕ): {phi_deg:.2f}°")

                doc.add_heading("Procedure", level=1)
                doc.add_paragraph(
                    "1. Place soil sample in shear box.\n"
                    "2. Apply normal load.\n"
                    "3. Measure horizontal deformation and proving ring readings.\n"
                    "4. Calculate shear stress and deformation.\n"
                    "5. Repeat for different normal stresses.\n"
                    "6. Determine cohesion and friction angle."
                )

                doc.add_heading("Formulas Used", level=1)
                doc.add_paragraph(
                    "Shear Force: Fs = Proving Ring Reading (div) * Proving Ring Constant (kg/div)\n"
                    "Shear Stress: tau = Fs / A (A = shear area in cm²)\n"
                    "Deformation: delta = Horizontal Deformation (div) * Dial Gauge Least Count (mm/div)\n"
                    "Mohr-Coulomb Failure: tau = c + sigma_n * tan(phi)"
                )

                # Add overall Mohr-Coulomb plot
                buf_plot = BytesIO()
                fig2.savefig(buf_plot, format="png", dpi=300, bbox_inches="tight")
                buf_plot.seek(0)
                doc.add_picture(buf_plot, width=Inches(6))
                doc.add_page_break()

                # Add each trial table and plot
                for t_idx, df_trial in enumerate(all_dfs):
                    doc.add_heading(f"Trial {t_idx+1} (σₙ={normal_stresses[t_idx]:.2f} kg/cm²)", level=2)
                    table = doc.add_table(rows=1, cols=len(df_trial.columns))
                    hdr_cells = table.rows[0].cells
                    for i_col, col_name in enumerate(df_trial.columns):
                        hdr_cells[i_col].text = col_name
                    for r in range(len(df_trial)):
                        row_cells = table.add_row().cells
                        for c_idx, col_name in enumerate(df_trial.columns):
                            val = df_trial[col_name].iloc[r]
                            row_cells[c_idx].text = f"{val:.3f}" if isinstance(val, (int,float)) else str(val)

                    # Trial plot
                    buf_trial = BytesIO()
                    fig_trial, ax_trial = plt.subplots()
                    ax_trial.plot(df_trial["Deformation (mm)"], df_trial["Shear Stress (kg/cm²)"], marker='o')
                    ax_trial.set_xlabel("Deformation (mm)")
                    ax_trial.set_ylabel("Shear Stress (kg/cm²)")
                    ax_trial.set_title(f"Trial {t_idx+1} Shear Stress vs Deformation")
                    ax_trial.grid(True)
                    fig_trial.savefig(buf_trial, format="png", dpi=300, bbox_inches="tight")
                    plt.close(fig_trial)
                    buf_trial.seek(0)
                    doc.add_picture(buf_trial, width=Inches(6))
                    doc.add_page_break()

                # Download Word Report
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    "📥 Download Direct Shear Test Report",
                    data=buffer.getvalue(),
                    file_name="Direct_Shear_Test_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("At least 2 trials with valid data are required to compute Cohesion and Friction Angle.")

    return None