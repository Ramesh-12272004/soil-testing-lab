import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO, StringIO

# -------------------------------------------------------
# PAGE CONFIG
# -------------------------------------------------------
st.set_page_config(
    page_title="Soil Tests Analysis",
    layout="wide",
    initial_sidebar_state="expanded"
)

# -------------------------------------------------------
# PREMIUM MODERN UI (BLUE PROFESSIONAL THEME)
# -------------------------------------------------------
st.markdown("""
    <style>

    /* App Background */
    .stApp {
        background: linear-gradient(135deg, #e8f1ff 0%, #cfe3ff 100%);
        font-family: "Segoe UI";
    }

    /* Sidebar Style */
    [data-testid="stSidebar"] {
        background: #0a4d94;
        padding-top: 25px;
    }
    [data-testid="stSidebar"] h2, 
    [data-testid="stSidebar"] h3, 
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span {
        color: white !important;
    }

    /* Sidebar Category Titles */
    .sidebar-title {
        color: #cfe8ff;
        font-size: 17px;
        font-weight: 600;
        padding-top: 15px;
    }

    /* Header Banner */
    .header-banner {
        background: linear-gradient(90deg, #0a68cc 0%, #0a4d94 100%);
        padding: 18px;
        border-radius: 12px;
        text-align: center;
        color: white;
        margin-bottom: 20px;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.2);
    }

    /* Start Button */
    .stButton > button {
        background-color: #0a68cc !important;
        color: white !important;
        border-radius: 10px !important;
        padding: 0.6em 1.2em !important;
        font-size: 18px !important;
        transition: 0.2s ease !important;
        border: none;
    }
    .stButton > button:hover {
        transform: scale(1.05);
        background-color: #084f99 !important;
    }

    /* Selectbox */
    .stSelectbox > div {
        border-radius: 10px !important;
        border: 2px solid #0a68cc !important;
        background: white !important;
        padding: 5px !important;
    }

    /* Result Card */
    .result-card {
        background: white;
        padding: 20px;
        border-radius: 12px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        animation: fadeIn 0.6s ease-in-out;
    }

    /* Fade Animation */
    @keyframes fadeIn {
        from {opacity: 0; transform: translateY(10px);}
        to {opacity: 1; transform: translateY(0);}
    }

    /* Report Buttons */
    .download-btn button {
        background-color: #28a745 !important;
        color: white !important;
        border-radius: 10px !important;
        padding: 0.6em 1.2em !important;
        font-size: 17px !important;
        transition: 0.25s !important;
    }
    .download-btn button:hover {
        transform: scale(1.05);
        background-color: #1e7e34 !important;
    }

    </style>
""", unsafe_allow_html=True)

# -------------------------------------------------------
# IMPORT ALL TEST MODULES
# -------------------------------------------------------
from tabs import (
    sieve_analysis,
    liquid_limit_casagrande,
    liquid_limit_cone,
    plastic_limit,
    core_cutter,
    specific_gravity,
    constant_head,
    variable_head,
    light_compaction,
    direct_shear,
    ucs_test,
    consolidation,
    cbr_test,
    vane_shear,
    triaxial_test
)

# -------------------------------------------------------
# SIDEBAR CATEGORIZED NAVIGATION
# -------------------------------------------------------
st.sidebar.markdown("<h2>🧪 Soil Tests Menu</h2>", unsafe_allow_html=True)

categories = {
    "🧪 Atterberg Limits": {
        "Liquid Limit (Casagrande)": liquid_limit_casagrande,
        "Liquid Limit (Cone Penetrometer)": liquid_limit_cone,
        "Plastic Limit": plastic_limit
    },

    "🌊 Permeability Tests": {
        "Constant Head Permeability": constant_head,
        "Variable Head Permeability": variable_head
    },

    "⚒️ Soil Properties": {
        "Sieve Analysis": sieve_analysis,
        "Specific Gravity": specific_gravity,
        "Core Cutter Method": core_cutter,
        "Light Compaction Test": light_compaction
    },

    "📊 Strength Tests": {
        "Unconfined Compressive Strength (UCS)": ucs_test,
        "Direct Shear Test": direct_shear,
        "Undrained Triaxial Test": triaxial_test,
        "Vane Shear Test": vane_shear
    },

    "🛣️ CBR & Consolidation": {
        "California Bearing Ratio (CBR) Test": cbr_test,
        "Consolidation Test": consolidation
    }
}

st.sidebar.markdown("---")

# Display categories
for cat, tests in categories.items():
    st.sidebar.markdown(f"<p class='sidebar-title'>{cat}</p>", unsafe_allow_html=True)
    for test_name in tests.keys():
        if st.sidebar.button(test_name):
            st.session_state['selected_test'] = test_name

# Default selected test
if 'selected_test' not in st.session_state:
    st.session_state['selected_test'] = "Sieve Analysis"

selected_test = st.session_state['selected_test']
selected_function = None

# Find the test function
for test_group in categories.values():
    if selected_test in test_group:
        selected_function = test_group[selected_test]

# -------------------------------------------------------
# HEADER BANNER
# -------------------------------------------------------
st.markdown("""
    <div class="header-banner">
        <h2>🧪 Soil Tests Analysis Dashboard</h2>
        <p>Analyze, visualize, and generate Soil Test Reports</p>
    </div>
""", unsafe_allow_html=True)

# -------------------------------------------------------
# RUN THE SELECTED TEST
# -------------------------------------------------------
returned_result = selected_function.run()

if returned_result is not None:
    st.markdown("<div class='result-card'>", unsafe_allow_html=True)
    if "completed_tests" not in st.session_state:
        st.session_state.completed_tests = {}
    st.session_state.completed_tests[selected_test] = returned_result
    st.markdown("</div>", unsafe_allow_html=True)

# -------------------------------------------------------
# REPORT GENERATION
# -------------------------------------------------------
if "completed_tests" in st.session_state and st.session_state.completed_tests:
    st.markdown("---")
    st.markdown("### 📄 Generate Combined Report")
    report_format = st.radio("Choose report format:", ("Excel", "Word (DOCX)"))

    if st.button("📥 Generate Combined Report"):
        
        # ----------------- Excel Report -----------------
        if report_format == "Excel":
            output = BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                for test_name, data in st.session_state.completed_tests.items():
                    if isinstance(data, dict):
                        for key, value in data.items():
                            if isinstance(value, pd.DataFrame):
                                sheet = f"{test_name[:20]}_{key[:10]}".replace(" ", "_")
                                value.to_excel(writer, sheet_name=sheet, index=False)
                    elif isinstance(data, pd.DataFrame):
                        sheet = test_name.replace(" ", "_")[:31]
                        data.to_excel(writer, sheet_name=sheet, index=False)

            st.download_button(
                "📥 Download Excel Report",
                data=output.getvalue(),
                file_name="Soil_Test_Report_Combined.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="excel_download",
            )

        # ----------------- Word Report -----------------
        else:
            doc = Document()
            doc.add_heading("🧪 Soil Tests Analysis Report", level=0)

            for test_name, test_data in st.session_state.completed_tests.items():
                doc.add_page_break()
                doc.add_heading(test_name, level=1)

                if isinstance(test_data, dict):
                    for sec_name, sec_data in test_data.items():
                        if isinstance(sec_data, pd.DataFrame):
                            doc.add_heading(sec_name, level=2)
                            table = doc.add_table(rows=1, cols=len(sec_data.columns))
                            hdr = table.rows[0].cells
                            for i, col in enumerate(sec_data.columns):
                                hdr[i].text = str(col)
                            for _, row in sec_data.iterrows():
                                row_cells = table.add_row().cells
                                for i, val in enumerate(row):
                                    row_cells[i].text = str(val)
                        else:
                            doc.add_paragraph(f"{sec_name}: {sec_data}")

                elif isinstance(test_data, pd.DataFrame):
                    doc.add_heading("Results Data", level=2)
                    table = doc.add_table(rows=1, cols=len(test_data.columns))
                    hdr = table.rows[0].cells
                    for i, col in enumerate(test_data.columns):
                        hdr[i].text = str(col)
                    for _, row in test_data.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)

            buffer = BytesIO()
            doc.save(buffer)

            st.download_button(
                "📥 Download Word Report",
                data=buffer.getvalue(),
                file_name="Soil_Test_Report_Combined.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="word_download",
            )
      
